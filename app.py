import streamlit as st
import pandas as pd
import plotly.graph_objects as go
import plotly.express as px
import requests
from docx import Document
from docx.shared import Inches
from io import BytesIO
import matplotlib.pyplot as plt

# --- 1. CONFIGURACIÓN INICIAL ---
APPSCRIPT_URL = st.secrets.get("https://script.google.com/macros/s/AKfycbz0TH94cn5tJWSJUjY-e0xALF2VlrzakwXo_hyzxy7TSGDKE3QW7OQOZSoFiCyNezniAQ/exec", None)

try:
    from st_gsheets_connection import GSheetsConnection
except ImportError:
    from streamlit_gsheets import GSheetsConnection

st.set_page_config(page_title="Consultoría BI Pro", layout="wide")

# --- 2. FUNCIONES DE DATOS ---
def get_data():
    try:
        conn = st.connection("gsheets", type=GSheetsConnection)
        df = conn.read(ttl=0)
        if df is not None and not df.empty:
            df.columns = df.columns.str.strip()
            # Asegurar tipos de datos numéricos
            df['Peso'] = pd.to_numeric(df['Peso'], errors='coerce').fillna(0)
            df['Calificacion'] = pd.to_numeric(df['Calificacion'], errors='coerce').fillna(0)
            # Asegurar que existan los Accionables
            if 'Accionables' not in df.columns:
                df['Accionables'] = ""
            return df
        return pd.DataFrame()
    except Exception as e:
        st.error(f"Error al conectar con Sheets: {e}")
        return pd.DataFrame()

if 'logos' not in st.session_state:
    st.session_state['logos'] = {}

# --- 3. GENERADOR DE REPORTE WORD ---
def generate_report_docx(empresa, nicho, df_n, prods_sel, resumen_df):
    doc = Document()
    if empresa in st.session_state['logos']:
        img_data = BytesIO(st.session_state['logos'][empresa])
        doc.add_picture(img_data, width=Inches(1.2))
    
    doc.add_heading(f'Informe Estratégico: {empresa}', 0)
    doc.add_paragraph(f'Análisis de Nicho: {nicho}')

    # Gráfico Radar para el documento
    plt.figure(figsize=(6, 6))
    ax = plt.subplot(111, polar=True)
    for p in prods_sel:
        d_p = df_n[df_n['Producto'] == p]
        if not d_p.empty:
            val = d_p['Calificacion'].tolist()
            val += val[:1]
            angles = [n / float(len(d_p['Factor'])) * 2 * 3.14159 for n in range(len(d_p['Factor']))]
            angles += angles[:1]
            ax.plot(angles, val, linewidth=1, label=p)
            ax.fill(angles, val, alpha=0.1)
            plt.xticks(angles[:-1], d_p['Factor'].tolist())
    plt.legend(loc='upper right', bbox_to_anchor=(0.1, 0.1))
    
    tmp_img = BytesIO()
    plt.savefig(tmp_img, format='png')
    doc.add_picture(tmp_img, width=Inches(4.5))
    plt.close()

    # Sección de Conclusiones y Accionables
    doc.add_heading('Plan de Acción por Producto', level=1)
    for p in prods_sel:
        score = resumen_df[resumen_df['Producto'] == p]['Puntaje Final'].values[0]
        doc.add_heading(f'Producto: {p} (Puntaje: {score:.2f})', level=2)
        
        acciones = df_n[df_n['Producto'] == p]['Accionables'].unique()
        for acc in acciones:
            if str(acc).strip() and str(acc) != "nan":
                doc.add_paragraph(f"{acc}", style='List Bullet')

    out = BytesIO()
    doc.save(out)
    return out.getvalue()

# --- 4. LÓGICA DE INTERFAZ ---
df = get_data()
if df.empty:
    st.info("Conectando con la base de datos...")
    st.stop()

# Sidebar: Filtros Globales y Logo
with st.sidebar:
    st.header("🏢 Filtros y Logo")
    emp_list = sorted(df['Empresa'].unique())
    emp_v = st.selectbox("Seleccionar Empresa Cliente", emp_list)
    
    logo_up = st.file_uploader("Subir Logo del Cliente", type=['png','jpg'])
    if logo_up:
        st.session_state['logos'][emp_v] = logo_up.getvalue()
    if emp_v in st.session_state['logos']:
        st.image(st.session_state['logos'][emp_v], width=120)
    
    st.divider()
    nicho_list = sorted(df[df['Empresa']==emp_v]['Nicho'].unique())
    nic_v = st.selectbox("Nicho Analizado", nicho_list)
    df_n = df[(df['Empresa']==emp_v) & (df['Nicho']==nic_v)].copy()
    
    prod_list = sorted(df_n['Producto'].unique())
    prods_v = st.multiselect("Productos a Comparar", prod_list, default=prod_list)

# Estructura de Pestañas
t_admin, t_radar, t_matriz = st.tabs(["⚙️ Gestión de Datos", "📊 Gráfico Radar", "📋 Matriz & Accionables"])

# Cálculos Globales
ajuste = 100 if df_n['Peso'].max() > 1.1 else 1
df_n['Ponderado'] = df_n['Calificacion'] * (df_n['Peso'] / ajuste)
res_df = df_n[df_n['Producto'].isin(prods_v)].groupby('Producto')['Ponderado'].sum().reset_index()
res_df.columns = ['Producto', 'Puntaje Final']

# --- PESTAÑA 1: GESTIÓN DE DATOS ---
with t_admin:
    st.header("🛠️ Registro de Factores, Calificaciones y Acciones")
    with st.form("form_registro", clear_on_submit=True):
        col1, col2 = st.columns(2)
        with col1:
            e_s = st.selectbox("Empresa", ["➕ Crear Nueva..."] + list(df['Empresa'].unique()))
            e_final = st.text_input("Nombre Nueva Empresa") if e_s == "➕ Crear Nueva..." else e_s
            
            n_opts = ["➕ Crear Nuevo..."] + (list(df[df['Empresa']==e_s]['Nicho'].unique()) if e_s != "➕ Crear Nueva..." else [])
            n_s = st.selectbox("Nicho", n_opts)
            n_final = st.text_input("Nombre Nuevo Nicho") if n_s == "➕ Crear Nuevo..." else n_s

        with col2:
            p_opts = ["➕ Crear Nuevo..."] + (list(df[(df['Empresa']==e_s)&(df['Nicho']==n_s)]['Producto'].unique()) if n_s != "➕ Crear Nuevo..." else [])
            p_s = st.selectbox("Producto", p_opts)
            p_final = st.text_input("Nombre Nuevo Producto") if p_s == "➕ Crear Nuevo..." else p_s
            
            f_opts = ["➕ Crear Nuevo Factor..."] + (list(df[(df['Empresa']==e_s)&(df['Nicho']==n_s)&(df['Producto']==p_s)]['Factor'].unique()) if p_s != "➕ Crear Nuevo..." else [])
            f_s = st.selectbox("Factor/Atributo", f_opts)
            f_final = st.text_input("Nombre Nuevo Factor") if f_s == "➕ Crear Nuevo Factor..." else f_s
            
            sub_col1, sub_col2 = st.columns(2)
            w_new = sub_col1.number_input("Peso (%)", 1, 100, 20)
            c_new = sub_col2.slider("Calificación (1-5)", 1, 5, 3)
            
        acc_new = st.text_area("🎯 Accionables / Recomendaciones Estratégicas")
        
        if st.form_submit_button("🚀 Guardar en Base de Datos"):
            if not APPSCRIPT_URL:
                st.error("Error: URL de AppScript no configurada en Secrets.")
            elif e_final and n_final and p_final and f_final:
                payload = {
                    "empresa": e_final, "nicho": n_final, "producto": p_final, 
                    "factor": f_final, "peso": w_new, "calificacion": c_new, 
                    "accionables": acc_new
                }
                try:
                    res = requests.post(APPSCRIPT_URL, json=payload, timeout=10)
                    if res.status_code == 200:
                        st.success(f"Guardado exitoso: {f_final}")
                        st.cache_data.clear()
                    else: st.error("Error al comunicar con Google Sheets.")
                except Exception as e: st.error(f"Fallo de conexión: {e}")
            else:
                st.warning("Completa todos los campos obligatorios.")

# --- PESTAÑA 2: GRÁFICO RADAR ---
with t_radar:
    if prods_v:
        fig = go.Figure()
        for p in prods_v:
            d_p = df_n[df_n['Producto']==p]
            fig.add_trace(go.Scatterpolar(r=d_p['Calificacion'], theta=d_p['Factor'], fill='toself', name=p))
        fig.update_layout(polar=dict(radialaxis=dict(visible=True, range=[0, 5])), title=f"Análisis Competitivo: {nic_v}")
        st.plotly_chart(fig, use_container_width=True)
    else:
        st.info("Selecciona productos en la barra lateral para ver el gráfico radar.")

# --- PESTAÑA 3: MATRIZ & ACCIONABLES ---
with t_matriz:
    st.header(f"Resumen de Resultados: {nic_v}")
    c_mat1, c_mat2 = st.columns([1, 1])
    
    with c_mat1:
        st.subheader("Puntajes Finales Ponderados")
        st.dataframe(res_df.style.background_gradient(cmap='Greens').format({'Puntaje Final': "{:.2f}"}), use_container_width=True)
        st.download_button("📂 Descargar Reporte Word", 
                           generate_report_docx(emp_v, nic_v, df_n, prods_v, res_df),
                           f"Reporte_{emp_v}_{nic_v}.docx")

    with c_mat2:
        st.subheader("Promedio de Desempeño por Factor")
        prom_factores = df_n[df_n['Producto'].isin(prods_v)].groupby('Factor')['Calificacion'].mean().reset_index()
        fig_bar = px.bar(prom_factores, x='Factor', y='Calificacion', color='Calificacion', range_y=[0,5], color_continuous_scale='RdYlGn')
        st.plotly_chart(fig_bar, use_container_width=True)

    st.divider()
    st.subheader("🎯 Plan de Acción por Producto")
    for p in prods_v:
        with st.expander(f"Ver Acciones para: {p}"):
            acciones = df_n[df_n['Producto'] == p]['Accionables'].unique()
            for a in acciones:
                if str(a).strip() and str(a) != "nan":
                    st.info(a)
                else:
                    st.write("No hay accionables registrados para este producto.")